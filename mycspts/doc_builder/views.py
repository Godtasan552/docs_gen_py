import os
import re
import json
import uuid
import subprocess
from django.shortcuts import render
from django.http import JsonResponse
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from docx import Document

LIBREOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"

def extract_fields_from_docx(filepath):
    """
    Index-based extraction to separate duplicate labels and prevent cross-talk.
    Returns a list of dictionaries containing label, unique ID, and paragraph index.
    """
    doc = Document(filepath)
    fields = []
    dot_pattern = r"[…\._]{2,}"
    
    label_counts = {}

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        label = None
        
        if re.search(dot_pattern, text):
            # Case 1: Label on same line
            match = re.search(r"([ก-๙a-zA-Z0-9\s(){}\[\]\-:]+?)\s*" + dot_pattern, text)
            if match:
                label = match.group(1).strip().strip(':').strip()
            # Case 2: Label on line above
            elif i > 0:
                prev_text = doc.paragraphs[i-1].text.strip().strip(':').strip()
                if prev_text and not re.search(dot_pattern, prev_text) and len(prev_text) < 100:
                    label = prev_text
            
            if label:
                # Count occurrences to handle duplicates (e.g., student 1, 2, 3)
                label_counts[label] = label_counts.get(label, 0) + 1
                display_label = label if label_counts[label] == 1 else f"{label} ({label_counts[label]})"
                
                fields.append({
                    'id': f"para_{i}_{label}", # Unique ID binding field to specific paragraph
                    'label': display_label,
                    'p_index': i
                })
                        
    return fields

def fill_docx(template_path, data, output_path):
    """
    Fills document using unique IDs to target specific paragraphs.
    Data key format: para_{index}_{label}
    """
    doc = Document(template_path)
    dot_pattern = r"[…\._]{2,}"
    
    for field_id, value in data.items():
        if field_id.startswith('para_'):
            try:
                # Extract paragraph index from ID
                parts = field_id.split('_')
                p_index = int(parts[1])
                
                if p_index < len(doc.paragraphs):
                    p = doc.paragraphs[p_index]
                    # Replace all dot patterns in this specific paragraph
                    for run in p.runs:
                        if re.search(dot_pattern, run.text):
                            run.text = re.sub(dot_pattern, value, run.text)
            except Exception as e:
                print(f"Fill Error for {field_id}: {e}")
    
    doc.save(output_path)

def convert_to_pdf(docx_path, output_dir):
    try:
        command = [
            LIBREOFFICE_PATH,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_path
        ]
        subprocess.run(command, check=True)
        return docx_path.replace('.docx', '.pdf')
    except Exception as e:
        print(f"PDF Conversion Error: {e}")
        return None

def index(request):
    return render(request, 'doc_builder/index.html')

def upload_file(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'uploads')
        if not os.path.exists(upload_dir):
            os.makedirs(upload_dir)
            
        fs = FileSystemStorage(location=upload_dir)
        filename = fs.save(f"{uuid.uuid4()}_{file.name}", file)
        filepath = fs.path(filename)
        
        fields = extract_fields_from_docx(filepath)
        return JsonResponse({
            'filename': filename,
            'fields': fields
        })
    return JsonResponse({'error': 'No file uploaded'}, status=400)

def generate(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            filename = data.get('filename')
            form_data = data.get('formData')
            
            input_path = os.path.join(settings.MEDIA_ROOT, 'uploads', filename)
            output_dir = os.path.join(settings.MEDIA_ROOT, 'output')
            
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            output_docx_name = f"filled_{filename}"
            output_docx_path = os.path.join(output_dir, output_docx_name)
            output_json_path = output_docx_path.replace('.docx', '.json')
            
            fill_docx(input_path, form_data, output_docx_path)
            
            with open(output_json_path, 'w', encoding='utf-8') as f:
                json.dump(form_data, f, ensure_ascii=False, indent=4)
                
            pdf_path = convert_to_pdf(output_docx_path, output_dir)
            
            if pdf_path and os.path.exists(pdf_path):
                return JsonResponse({
                    'success': True,
                    'pdf_url': f"{settings.MEDIA_URL}output/{os.path.basename(pdf_path)}",
                    'json_url': f"{settings.MEDIA_URL}output/{os.path.basename(output_json_path)}"
                })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Failed'}, status=500)
